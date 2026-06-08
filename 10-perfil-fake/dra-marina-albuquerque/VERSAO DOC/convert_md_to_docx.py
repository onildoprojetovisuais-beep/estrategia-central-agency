import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_heading_style(para, level, text):
    """Apply heading formatting"""
    from docx.shared import Pt, RGBColor
    run = para.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x3c, 0x78)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

def add_cover_page(doc, title):
    """Add a simple cover page"""
    # Cover title
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_para.paragraph_format.space_before = Pt(120)
    cover_run = cover_para.add_run(title.upper())
    cover_run.font.size = Pt(28)
    cover_run.font.bold = True
    cover_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Subtitle line
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Dra. Marina Albuquerque")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_run.font.italic = True

    # Separator
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("─" * 40)
    sep_run.font.color.rgb = RGBColor(0x0f, 0x3c, 0x78)

    doc.add_page_break()

def parse_inline(text):
    """Parse inline markdown: **bold**, *italic*, `code`"""
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append(('normal', text[last:m.start()]))
        if m.group(2):
            segments.append(('bold', m.group(2)))
        elif m.group(3):
            segments.append(('italic', m.group(3)))
        elif m.group(4):
            segments.append(('code', m.group(4)))
        last = m.end()
    if last < len(text):
        segments.append(('normal', text[last:]))
    return segments

def add_inline_runs(para, text):
    segments = parse_inline(text)
    for style, content in segments:
        run = para.add_run(content)
        if style == 'bold':
            run.font.bold = True
        elif style == 'italic':
            run.font.italic = True
        elif style == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

def process_table(doc, table_lines):
    """Process markdown table"""
    rows = []
    for line in table_lines:
        if re.match(r'\|[\s\-:]+\|', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = row.cells[j]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(cell_text)
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
                    # Header background color
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '1a1a2e')
                    tcPr.append(shd)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Get title from first H1
    title = "Documento"
    for line in lines:
        if line.startswith('# '):
            title = line[2:].strip()
            break

    add_cover_page(doc, title)

    i = 0
    in_table = False
    table_buffer = []
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                if code_buffer:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    for code_line in code_buffer:
                        run = para.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Tables
        if line.startswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        else:
            if table_buffer:
                process_table(doc, table_buffer)
                table_buffer = []

        # Headings
        if line.startswith('### '):
            para = doc.add_paragraph()
            set_heading_style(para, 3, line[4:].strip())
        elif line.startswith('## '):
            para = doc.add_paragraph()
            set_heading_style(para, 2, line[3:].strip())
        elif line.startswith('# '):
            para = doc.add_paragraph()
            set_heading_style(para, 1, line[2:].strip())

        # Horizontal rule
        elif line.strip() in ['---', '***', '___']:
            para = doc.add_paragraph()
            run = para.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

        # Unordered lists
        elif re.match(r'^[\*\-\+] ', line):
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, line[2:].strip())

        elif re.match(r'^\s{2,4}[\*\-\+] ', line):
            para = doc.add_paragraph(style='List Bullet 2')
            para.paragraph_format.left_indent = Cm(2)
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, re.sub(r'^\s+[\*\-\+] ', '', line).strip())

        # Ordered lists
        elif re.match(r'^\d+\. ', line):
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, re.sub(r'^\d+\. ', '', line).strip())

        # Blockquote
        elif line.startswith('> '):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.5)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(line[2:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Empty line
        elif line.strip() == '':
            if i > 0 and lines[i-1].strip() != '':
                doc.add_paragraph()

        # Normal paragraph
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            add_inline_runs(para, line.strip())

        i += 1

    # Process any remaining table
    if table_buffer:
        process_table(doc, table_buffer)

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"CONVERTED: {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
