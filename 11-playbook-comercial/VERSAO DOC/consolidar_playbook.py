import os
import copy
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE_FOLDER = r"C:\Projetos ClaudeCode\PIXEL RAIN\pixel-rain-agency\11-playbook-comercial\VERSAO DOC"
OUTPUT_FOLDER = os.path.join(SOURCE_FOLDER, "PLAYBOOK CONSOLIDADO")
OUTPUT_FILE = os.path.join(OUTPUT_FOLDER, "PLAYBOOK COMERCIAL PIXEL RAIN AGENCY.docx")

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Estrutura do documento final
STRUCTURE = [
    {
        "heading": "01 — FUNDAMENTOS",
        "docs": [
            ("00-glossario-e-indice.docx", "Glossário e Índice"),
            ("01-icp-definitivo.docx", "ICP Definitivo"),
            ("01-auditoria-conformidade-jornada.docx", "Auditoria de Conformidade da Jornada"),
            ("02-segmentos-prioritarios.docx", "Segmentos Prioritários"),
            ("03-criterios-qualificacao.docx", "Critérios de Qualificação"),
        ]
    },
    {
        "heading": "02 — PROSPECÇÃO",
        "docs": [
            ("04-processo-prospeccao.docx", "Processo de Prospecção"),
            ("05-funil-comercial.docx", "Funil Comercial"),
            ("06-cadencia-contatos.docx", "Cadência de Contatos"),
            ("10-spin-selling-adaptado.docx", "Social Selling / SPIN Selling Adaptado"),
        ]
    },
    {
        "heading": "03 — SCRIPTS",
        "docs": [
            ("07-scripts-ligacao.docx", "Scripts de Ligação"),
            ("08-scripts-whatsapp.docx", "Scripts WhatsApp"),
            ("09-scripts-email.docx", "Scripts de E-mail"),
        ]
    },
    {
        "heading": "04 — DIAGNÓSTICO",
        "docs": [
            ("11-diagnostico-comercial.docx", "Diagnóstico Comercial"),
        ]
    },
    {
        "heading": "05 — FOLLOW-UP",
        "docs": [
            ("12-processo-follow-up.docx", "Processo de Follow-up"),
        ]
    },
    {
        "heading": "06 — FECHAMENTO",
        "docs": [
            ("13-processo-fechamento.docx", "Processo de Fechamento"),
        ]
    },
    {
        "heading": "07 — GESTÃO COMERCIAL",
        "docs": [
            ("14-metricas.docx", "Métricas Comerciais"),
            ("15-rotina-diaria.docx", "Rotina Diária"),
            ("16-rotina-semanal.docx", "Rotina Semanal"),
            ("17-rotina-mensal.docx", "Rotina Mensal"),
            ("18-checklist-operacional.docx", "Checklist Operacional"),
            ("19-rotina-execucao-nil.docx", "Rotina de Execução do Nil"),
            ("20-fase-zero-implantacao.docx", "Fase Zero — Implantação"),
        ]
    },
]


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def add_toc(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar_begin, instrText, fldChar_sep, fldChar_end])

    note = doc.add_paragraph()
    note_run = note.add_run("[ Abra no Word e pressione Ctrl+A, depois F9 para atualizar o Sumário ]")
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    note_run.font.size = Pt(9)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def copy_images(src_doc, dst_doc, errors, src_name):
    """Copy image parts from src_doc to dst_doc, return rId mapping."""
    img_rId_map = {}
    try:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        import posixpath

        for rel in src_doc.part.rels.values():
            if 'image' in rel.reltype.lower():
                try:
                    img_part = rel.target_part
                    orig_name = posixpath.basename(str(img_part.partname))
                    new_partname = PackURI('/word/media/' + orig_name)
                    new_img_part = Part(new_partname, img_part.content_type, img_part.blob)
                    new_rId = dst_doc.part.relate_to(new_img_part, rel.reltype)
                    img_rId_map[rel.rId] = new_rId
                except Exception as e:
                    errors.append(f"  AVISO imagem em {src_name}: {e}")
    except Exception as e:
        errors.append(f"  AVISO ao processar imagens de {src_name}: {e}")
    return img_rId_map


def remap_image_rids(element, img_rId_map):
    """Walk element tree and remap image relationship IDs."""
    if not img_rId_map:
        return
    a_blip = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    r_embed = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    for blip in element.iter(a_blip):
        old_rId = blip.get(r_embed)
        if old_rId and old_rId in img_rId_map:
            blip.set(r_embed, img_rId_map[old_rId])


def copy_content_from_doc(src_path, dst_doc, errors):
    """Copy all body content from a source .docx into dst_doc."""
    src_name = os.path.basename(src_path)
    try:
        src_doc = Document(src_path)
    except Exception as e:
        errors.append(f"ERRO ao abrir {src_name}: {e}")
        return False

    img_rId_map = copy_images(src_doc, dst_doc, errors, src_name)

    src_body = src_doc.element.body
    dst_body = dst_doc.element.body

    # Find the final sectPr in dst to insert before it
    dst_sectPr = dst_body.find(qn('w:sectPr'))

    copied = 0
    for element in src_body.iterchildren():
        # Skip section properties (page setup) of source
        if element.tag == qn('w:sectPr'):
            continue
        new_element = copy.deepcopy(element)
        remap_image_rids(new_element, img_rId_map)
        if dst_sectPr is not None:
            dst_sectPr.addprevious(new_element)
        else:
            dst_body.append(new_element)
        copied += 1

    return True


# ─────────────────────────────────────────────
print("=" * 60)
print("CONSOLIDADOR — PLAYBOOK COMERCIAL PIXEL RAIN AGENCY")
print("=" * 60)

dst_doc = Document()

# ── CAPA ──────────────────────────────────────
cover_title = dst_doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("\n\n\n\nPLAYBOOK COMERCIAL")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

agency = dst_doc.add_paragraph()
agency.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = agency.add_run("PIXEL RAIN AGENCY")
run2.bold = True
run2.font.size = Pt(24)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

version = dst_doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = version.add_run("\nVersão Consolidada")
run3.italic = True
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ── SUMÁRIO ───────────────────────────────────
add_page_break(dst_doc)
sumario_heading = dst_doc.add_heading("SUMÁRIO", level=1)
sumario_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_toc(dst_doc)

# ── CAPÍTULOS ─────────────────────────────────
incorporated = []
not_found = []
errors = []

all_docs_in_structure = []
for ch in STRUCTURE:
    for fn, _ in ch["docs"]:
        all_docs_in_structure.append(fn)

for chapter in STRUCTURE:
    add_page_break(dst_doc)
    ch_heading = dst_doc.add_heading(chapter["heading"], level=1)
    print(f"\n  [{chapter['heading']}]")

    for filename, display_name in chapter["docs"]:
        src_path = os.path.join(SOURCE_FOLDER, filename)

        if not os.path.exists(src_path):
            not_found.append(filename)
            errors.append(f"NAO ENCONTRADO: {filename}")
            print(f"    [X] NAO ENCONTRADO: {filename}")
            continue

        add_page_break(dst_doc)
        dst_doc.add_heading(display_name, level=2)

        success = copy_content_from_doc(src_path, dst_doc, errors)
        if success:
            incorporated.append(filename)
            print(f"    [OK] {filename}")
        else:
            errors.append(f"FALHA ao incorporar: {filename}")
            print(f"    [FALHA] {filename}")

# ── VERIFICAR ARQUIVOS NÃO MAPEADOS ───────────
all_docx = [f for f in os.listdir(SOURCE_FOLDER) if f.endswith('.docx')]
unmapped = [f for f in all_docx if f not in all_docs_in_structure]

# ── SALVAR ────────────────────────────────────
print(f"\nSalvando em: {OUTPUT_FILE}")
dst_doc.save(OUTPUT_FILE)

# ── RELATÓRIO FINAL ───────────────────────────
print("\n" + "=" * 60)
print("RELATÓRIO FINAL")
print("=" * 60)
print(f"Arquivos .docx encontrados na pasta : {len(all_docx)}")
print(f"Documentos incorporados             : {len(incorporated)}")
print(f"Documentos não encontrados          : {len(not_found)}")
if unmapped:
    print(f"Arquivos na pasta fora da estrutura : {len(unmapped)}")
    for f in unmapped:
        print(f"  - {f}")
print()
print("DOCUMENTOS INCORPORADOS:")
for i, f in enumerate(incorporated, 1):
    print(f"  {i:02d}. {f}")
if not_found:
    print("\nNÃO ENCONTRADOS:")
    for f in not_found:
        print(f"  - {f}")
if errors:
    print("\nAVISOS / ERROS:")
    for e in errors:
        print(f"  {e}")
print()
print(f"Arquivo salvo em:\n  {OUTPUT_FILE}")
print("=" * 60)
