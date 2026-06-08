#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Converte 12 arquivos .md para .docx com formatacao profissional.
Pixel Rain Agency - Processo Padrao N1
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────
BASE_DIR = r"C:\Projetos ClaudeCode\PIXEL RAIN\pixel-rain-agency\08-escada-de-valor\n1-entregaveis\processo-padrao"
OUTPUT_DIR = os.path.join(BASE_DIR, "VERSAO DOC")

FILES = [
    "E01-kickoff.md",
    "E02-checklist-coleta.md",
    "E03-auditoria-presenca-digital.md",
    "E04-otimizacao-google-business.md",
    "E05-painel-omnichannel.md",
    "E06-instalacao-rastreamento.md",
    "E07-landing-pages-funis.md",
    "E08-campanhas-trafego.md",
    "E09-calendario-editorial.md",
    "E10-operacao-conteudo.md",
    "E11-relatorio-30-dias.md",
    "E12-rotina-mensal.md",
]

# ─────────────────────────────────────────────
# COLORS
# ─────────────────────────────────────────────
COLOR_BRAND    = RGBColor(0x1A, 0x1A, 0x2E)  # Azul escuro (capa / H1)
COLOR_H1       = RGBColor(0x0F, 0x3C, 0x78)  # Azul medio
COLOR_H2       = RGBColor(0x16, 0x21, 0x3E)  # Azul escuro
COLOR_H3       = RGBColor(0x1A, 0x5,  0x7B)  # Azul
COLOR_H4       = RGBColor(0x2E, 0x86, 0xAB)  # Azul medio-claro
COLOR_TABLE_HDR = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_TABLE_ALT = RGBColor(0xF0, 0xF4, 0xF8)
COLOR_ACCENT   = RGBColor(0xE8, 0x4A, 0x5F)  # vermelho para separadores
COLOR_CODE_BG  = RGBColor(0xF5, 0xF5, 0xF0)
COLOR_NOTE_BG  = RGBColor(0xFE, 0xF9, 0xE7)


def set_cell_bg(cell, color_rgb):
    """Define cor de fundo de uma celula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(color_rgb[0], color_rgb[1], color_rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    """Aplica bordas visiveis em toda a tabela."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '1A1A2E')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def apply_inline_formatting(run_parent, text):
    """
    Processa texto com **negrito** e *italico* inline e adiciona runs ao paragrafo.
    run_parent: paragrafo (p) ao qual adicionar os runs.
    """
    # Pattern: **texto** ou *texto*
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        # texto antes do match
        before = text[last:m.start()]
        if before:
            run_parent.add_run(before)
        if m.group(2):  # **bold**
            r = run_parent.add_run(m.group(2))
            r.bold = True
        elif m.group(3):  # *italic*
            r = run_parent.add_run(m.group(3))
            r.italic = True
        elif m.group(4):  # `code`
            r = run_parent.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
        last = m.end()
    remainder = text[last:]
    if remainder:
        run_parent.add_run(remainder)


def add_cover_page(doc, title_text, subtitle_text):
    """Adiciona pagina de capa ao documento."""
    # Capa: fundo azul escuro simulado por texto com cor
    doc.add_paragraph()  # espaco inicial

    # Linha decorativa
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("━" * 45)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Logo text / Agency name
    agency = doc.add_paragraph()
    agency.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = agency.add_run("PIXEL RAIN AGENCY")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_H2
    r.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    # Titulo principal da capa
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = COLOR_H1
    r.font.name = 'Calibri'

    doc.add_paragraph()

    # Subtitulo
    if subtitle_text:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub_p.add_run(subtitle_text)
        r.font.size = Pt(14)
        r.font.color.rgb = COLOR_H4
        r.font.name = 'Calibri'
        r.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Linha decorativa inferior
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sep2.add_run("━" * 45)
    run2.font.color.rgb = COLOR_ACCENT
    run2.font.size = Pt(10)

    doc.add_paragraph()

    # Tag de classificacao
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("Processo Operacional Interno  |  N1 — Plano de Atração  |  Versão 2.0 — Maio 2026")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.font.name = 'Calibri'

    # Quebra de pagina apos a capa
    doc.add_page_break()


def set_paragraph_spacing(p, before=0, after=4, line_spacing=None):
    """Ajusta espacamento antes/depois do paragrafo."""
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line_spacing)


def add_heading(doc, text, level):
    """Adiciona titulo formatado."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else (10 if level == 2 else 6), after=4)

    if level == 1:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = COLOR_H1
        r.font.name = 'Calibri'
        # linha abaixo
        sep = doc.add_paragraph()
        set_paragraph_spacing(sep, before=0, after=6)
        rs = sep.add_run("─" * 60)
        rs.font.size = Pt(8)
        rs.font.color.rgb = COLOR_H1

    elif level == 2:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = COLOR_H2
        r.font.name = 'Calibri'

    elif level == 3:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = COLOR_H3
        r.font.name = 'Calibri'

    elif level == 4:
        r = p.add_run(text)
        r.bold = True
        r.italic = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = COLOR_H4
        r.font.name = 'Calibri'

    return p


def add_hr(doc):
    """Adiciona linha horizontal separadora."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6)
    r = p.add_run("─" * 80)
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_blockquote(doc, text):
    """Adiciona paragrafo de blockquote com barra lateral."""
    # Simular blockquote com indentacao e fonte ligeiramente diferente
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.2)
    set_paragraph_spacing(p, before=4, after=4)
    apply_inline_formatting(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.font.name = 'Calibri'
        run.italic = True
    return p


def add_code_block(doc, lines):
    """Adiciona bloco de codigo monoespaco."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        set_paragraph_spacing(p, before=0, after=1)
        r = p.add_run(line if line else " ")
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # espaco depois
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=4)


def parse_table_rows(lines, start_idx):
    """
    Coleta as linhas de uma tabela markdown a partir de start_idx.
    Retorna (rows, next_idx) onde rows e lista de listas de celulas.
    """
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # linha separadora (---|---)
        if re.match(r'^\|[\s\-\|:]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')]
        # remover primeiro e ultimo vazios
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """Adiciona tabela formatada ao documento."""
    if not rows:
        return
    # determinar numero maximo de colunas
    ncols = max(len(r) for r in rows)
    # garantir que todas as linhas tenham o mesmo numero de colunas
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        tr = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]
            # limpar conteudo
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            if row_idx == 0:
                # Cabecalho
                set_cell_bg(cell, (0x1A, 0x1A, 0x2E))
                apply_inline_formatting(p, cell_text)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = 'Calibri'
            else:
                # Linha alternada
                if row_idx % 2 == 0:
                    set_cell_bg(cell, (0xF0, 0xF4, 0xF8))
                apply_inline_formatting(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'

    set_table_border(table)

    # espaco depois da tabela
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=6)


def add_list_item(doc, text, ordered=False, number=None, indent=0):
    """Adiciona item de lista (bullet ou numerada)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    set_paragraph_spacing(p, before=1, after=1)

    if ordered:
        prefix = f"{number}. "
        r = p.add_run(prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
    else:
        prefix = "• " if indent == 0 else "◦ "
        r = p.add_run(prefix)
        r.font.size = Pt(10.5)
        r.font.color.rgb = COLOR_ACCENT
        r.font.name = 'Calibri'

    apply_inline_formatting(p, text)
    for run in p.runs:
        if not run.bold or (ordered and run.text == prefix):
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
    return p


def add_checkbox_item(doc, text, checked=False):
    """Adiciona item de lista com checkbox [ ] ou [x]."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    set_paragraph_spacing(p, before=1, after=1)

    box = "☑ " if checked else "☐ "
    r = p.add_run(box)
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_H4 if not checked else RGBColor(0x00, 0x80, 0x00)
    r.font.name = 'Calibri'

    apply_inline_formatting(p, text)
    for run in p.runs:
        if run.text != box:
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
    return p


def add_normal_paragraph(doc, text):
    """Adiciona paragrafo normal."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=5)
    apply_inline_formatting(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
    return p


# ─────────────────────────────────────────────
# PARSER PRINCIPAL
# ─────────────────────────────────────────────

def parse_and_build(doc, lines):
    """Processa as linhas do markdown e constroi o documento."""
    i = 0
    ordered_counters = {}  # indent_level -> counter

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip('\n')
        stripped = line.strip()

        # Linha vazia
        if not stripped:
            i += 1
            continue

        # Quebra de pagina (---) -> separador horizontal
        if re.match(r'^---+$', stripped):
            add_hr(doc)
            i += 1
            continue

        # Headings
        h4 = re.match(r'^####\s+(.*)', stripped)
        h3 = re.match(r'^###\s+(.*)', stripped)
        h2 = re.match(r'^##\s+(.*)', stripped)
        h1 = re.match(r'^#\s+(.*)', stripped)

        if h4:
            add_heading(doc, h4.group(1), 4)
            i += 1
            continue
        if h3:
            add_heading(doc, h3.group(1), 3)
            i += 1
            continue
        if h2:
            add_heading(doc, h2.group(1), 2)
            i += 1
            continue
        if h1:
            add_heading(doc, h1.group(1), 1)
            i += 1
            continue

        # Tabela markdown
        if stripped.startswith('|'):
            rows, next_i = parse_table_rows(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # Bloco de codigo (``` ... ```)
        if stripped.startswith('```'):
            # coletar ate o fechamento
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # pular o fechamento
            add_code_block(doc, code_lines)
            continue

        # Blockquote (> texto)
        if stripped.startswith('>'):
            # coletar linhas consecutivas de blockquote
            bq_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                bq_line = lines[i].strip()
                bq_line = re.sub(r'^>\s?', '', bq_line)
                bq_lines.append(bq_line)
                i += 1
            # unir em paragrafos (linhas vazias dentro do blockquote = nova secao)
            bq_text = ' '.join(bq_lines)
            add_blockquote(doc, bq_text)
            continue

        # Checkbox [ ] ou [x]
        cb_match = re.match(r'^[-*]\s+\[([xX ])\]\s+(.*)', stripped)
        if cb_match:
            checked = cb_match.group(1).lower() == 'x'
            text = cb_match.group(2)
            add_checkbox_item(doc, text, checked)
            i += 1
            continue

        # Lista nao-ordenada
        ul_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if ul_match:
            indent = len(ul_match.group(1)) // 2
            text = ul_match.group(3)
            add_list_item(doc, text, ordered=False, indent=indent)
            i += 1
            continue

        # Lista ordenada
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ol_match:
            indent = len(ol_match.group(1)) // 2
            num = int(ol_match.group(2))
            text = ol_match.group(3)
            add_list_item(doc, text, ordered=True, number=num, indent=indent)
            i += 1
            continue

        # Paragrafo normal
        add_normal_paragraph(doc, stripped)
        i += 1


def extract_cover_info(lines):
    """Extrai titulo e subtitulo das primeiras linhas do MD."""
    title = ""
    subtitle = ""
    for line in lines[:5]:
        stripped = line.strip()
        h1 = re.match(r'^#\s+(.*)', stripped)
        if h1 and not title:
            title = h1.group(1)
        elif stripped.startswith('**') and not subtitle:
            text = stripped.strip('*').strip()
            subtitle = text
    return title, subtitle


def convert_md_to_docx(md_path, docx_path):
    """Converte um arquivo .md para .docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()

    # Configurar margens da pagina
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # Estilo de fonte padrao do documento
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Extrair info para capa
    title, subtitle = extract_cover_info(lines)

    # Adicionar capa
    add_cover_page(doc, title, subtitle)

    # Processar o conteudo (pular o H1 inicial e o subtitulo bold — ja estao na capa)
    # Encontrar onde comecar (apos as linhas de titulo/subtitulo iniciais)
    start = 0
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r'^#\s+', stripped) or (stripped.startswith('**') and idx < 5):
            start = idx + 1
        else:
            break

    # Parsear o corpo do documento
    parse_and_build(doc, lines[start:])

    doc.save(docx_path)
    return True


# ─────────────────────────────────────────────
# EXECUCAO PRINCIPAL
# ─────────────────────────────────────────────

def main():
    # Criar pasta de destino
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"Pasta de saida: {OUTPUT_DIR}\n")

    results = []
    for md_file in FILES:
        md_path = os.path.join(BASE_DIR, md_file)
        docx_name = os.path.splitext(md_file)[0] + '.docx'
        docx_path = os.path.join(OUTPUT_DIR, docx_name)

        if not os.path.exists(md_path):
            print(f"  [ERRO] Arquivo nao encontrado: {md_file}")
            results.append((md_file, docx_name, False, 0))
            continue

        try:
            convert_md_to_docx(md_path, docx_path)
            size = os.path.getsize(docx_path)
            print(f"  [OK] {md_file} -> {docx_name} ({size:,} bytes)")
            results.append((md_file, docx_name, True, size))
        except Exception as e:
            print(f"  [ERRO] {md_file}: {e}")
            import traceback; traceback.print_exc()
            results.append((md_file, docx_name, False, 0))

    print("\n" + "=" * 60)
    print("RESUMO FINAL")
    print("=" * 60)
    ok = [r for r in results if r[2]]
    fail = [r for r in results if not r[2]]
    print(f"Convertidos com sucesso: {len(ok)}/12")
    if fail:
        print("Falhas:")
        for f in fail:
            print(f"  - {f[0]}")
    print("\nArquivos gerados em VERSAO DOC:")
    for r in ok:
        print(f"  {r[1]:45s}  {r[3]:>10,} bytes")


if __name__ == '__main__':
    main()
