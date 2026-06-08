#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor de .md para .docx com formatacao profissional.
Preserva 100% do conteudo original.
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


BASE_DIR = Path(r"C:\Projetos ClaudeCode\PIXEL RAIN\pixel-rain-agency\08-escada-de-valor\n1-entregaveis\projeto-dra-marina")
OUTPUT_DIR = BASE_DIR / "VERSAO DOC"

# Paleta de cores profissional
COLOR_H1 = RGBColor(0x1A, 0x37, 0x5E)      # Azul escuro profissional
COLOR_H2 = RGBColor(0x2E, 0x5E, 0x8B)      # Azul medio
COLOR_H3 = RGBColor(0x3A, 0x7D, 0xBD)      # Azul claro
COLOR_COVER_BG = RGBColor(0x1A, 0x37, 0x5E)
COLOR_CODE_BG = RGBColor(0xF2, 0xF4, 0xF7)
COLOR_RULE = RGBColor(0xCC, 0xD6, 0xE0)
COLOR_TABLE_HEADER = RGBColor(0x1A, 0x37, 0x5E)
COLOR_TABLE_STRIPE = RGBColor(0xF0, 0xF4, 0xF8)
COLOR_BLOCKQUOTE = RGBColor(0x2E, 0x5E, 0x8B)


def set_cell_bg(cell, color_rgb):
    """Define cor de fundo de celula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    # RGBColor e uma tuple (r, g, b) ou tem __str__ que retorna hex
    hex_color = str(color_rgb)  # retorna '1A375E' direto
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Adiciona linha horizontal decorativa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCD6E0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def configure_document_styles(doc):
    """Configura estilos base do documento."""
    # Estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Margem da pagina
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def add_cover_page(doc, title, subtitle_lines=None):
    """Adiciona pagina de capa profissional."""
    # Espaco superior
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Linha decorativa superior
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A375E')
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    # Titulo principal da capa
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_H1

    # Subtitulos da capa
    if subtitle_lines:
        for line in subtitle_lines:
            if line.strip():
                sub_para = doc.add_paragraph()
                sub_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_para.paragraph_format.space_before = Pt(3)
                sub_para.paragraph_format.space_after = Pt(3)
                run = sub_para.add_run(line.strip('*').strip())
                run.font.name = 'Calibri'
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    doc.add_paragraph()

    # Linha decorativa inferior
    rule_para2 = doc.add_paragraph()
    rule_para2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = rule_para2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '1A375E')
    pBdr2.append(top)
    pPr2.append(pBdr2)

    # Quebra de pagina apos capa
    doc.add_page_break()


def apply_inline_formatting(run_text, paragraph):
    """Aplica formatacao inline (negrito, italico, code) num paragrafo."""
    # Padroes de formatacao inline
    # Precisa ser processado em ordem: bold+italic, bold, italic, code
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|__(.+?)__)')

    pos = 0
    for match in pattern.finditer(run_text):
        # Texto antes do match
        before = run_text[pos:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        full = match.group(0)
        if full.startswith('***') and full.endswith('***'):
            run = paragraph.add_run(match.group(2))
            run.font.bold = True
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif full.startswith('**') and full.endswith('**'):
            run = paragraph.add_run(match.group(3))
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif full.startswith('`') and full.endswith('`'):
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif (full.startswith('*') and full.endswith('*')) or (full.startswith('__') and full.endswith('__')):
            inner = match.group(5) or match.group(6)
            run = paragraph.add_run(inner)
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        pos = match.end()

    # Texto restante
    remaining = run_text[pos:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def add_styled_paragraph(doc, text, level=0, list_type=None, list_number=None):
    """Adiciona paragrafo com formatacao inline aplicada."""
    p = doc.add_paragraph()

    if list_type == 'bullet':
        p.style = doc.styles['List Bullet']
        p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    elif list_type == 'numbered':
        p.style = doc.styles['List Number']
        p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    else:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(14)

    apply_inline_formatting(text, p)
    return p


def add_heading(doc, text, level):
    """Adiciona titulo com estilo personalizado."""
    p = doc.add_paragraph()

    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = COLOR_H1
        # Borda inferior
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A375E')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = COLOR_H2
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_H3
    elif level >= 4:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x66, 0x88)

    return p


def add_code_block(doc, lines):
    """Adiciona bloco de codigo com fundo cinza."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        # Fundo cinza claro via shading no paragrafo
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F4F7')
        pPr.append(shd)

        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)


def add_blockquote(doc, text):
    """Adiciona blockquote com barra lateral e recuo."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Borda esquerda azul
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2E5E8B')
    pBdr.append(left)
    pPr.append(pBdr)

    # Fundo azul muito claro
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF4FA')
    pPr.append(shd)

    apply_inline_formatting(text, p)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x33, 0x44, 0x55)


def parse_table(doc, table_lines):
    """Faz parse e renderiza uma tabela Markdown."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line or re.match(r'^\|?[\s\-:]+\|', line):
            continue
        cells = [c.strip() for c in re.split(r'(?<!\\)\|', line) if c.strip() != '']
        if cells:
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # Normaliza todas as linhas para o mesmo numero de colunas
    rows = [r + [''] * (max_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.autofit = True

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            if i == 0:
                # Cabecalho
                set_cell_bg(cell, COLOR_TABLE_HEADER)
                run = p.add_run(cell_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Linhas alternadas
                if i % 2 == 0:
                    set_cell_bg(cell, COLOR_TABLE_STRIPE)
                apply_inline_formatting(cell_text, p)
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

    # Espaco apos tabela
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Converte um arquivo .md para .docx com formatacao profissional."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.splitlines()

    doc = Document()
    configure_document_styles(doc)

    # Extrair titulo e subtitulos para a capa
    cover_title = md_path.stem  # fallback
    cover_subtitles = []

    for i, line in enumerate(lines[:10]):
        if line.startswith('# '):
            cover_title = line[2:].strip()
        elif line.startswith('**') and line.endswith('**') and not line.startswith('**N'):
            pass
        elif line.startswith('**') and line.endswith('**'):
            cover_subtitles.append(line.strip('*').strip())
        elif line.strip() and not line.startswith('#') and not line.startswith('**') and not line.startswith('---'):
            if len(cover_subtitles) < 3:
                cover_subtitles.append(line.strip('*').strip())

    add_cover_page(doc, cover_title, cover_subtitles[:4])

    # Processar o conteudo linha a linha
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Bloco de codigo
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                # Fechar bloco
                in_code_block = False
                add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tabela
        stripped = line.strip()
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [stripped]
            else:
                table_lines.append(stripped)
            i += 1
            # Verificar se a proxima linha tambem e tabela
            if i < len(lines) and not lines[i].strip().startswith('|'):
                parse_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
        else:
            if in_table:
                parse_table(doc, table_lines)
                in_table = False
                table_lines = []

        # Linha horizontal
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Titulos
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            bq_text = stripped[2:]
            # Coletar linhas de blockquote continuas
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                bq_text += ' ' + lines[i].strip()[2:]
            add_blockquote(doc, bq_text)
            i += 1
            continue

        # Lista com bullet (-, *, +)
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(3)
            add_styled_paragraph(doc, text, level=indent, list_type='bullet')
            i += 1
            continue

        # Lista numerada
        num_match = re.match(r'^(\s*)(\d+)[.)]\s+(.+)$', line)
        if num_match:
            indent = len(num_match.group(1)) // 2
            text = num_match.group(3)
            add_styled_paragraph(doc, text, level=indent, list_type='numbered')
            i += 1
            continue

        # Linha em branco
        if not stripped:
            i += 1
            continue

        # Paragrafo normal
        add_styled_paragraph(doc, stripped)
        i += 1

    # Fechar tabela pendente
    if in_table and table_lines:
        parse_table(doc, table_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main():
    print(f"Pasta base: {BASE_DIR}")
    print(f"Pasta de saida: {OUTPUT_DIR}")
    print()

    # Encontrar todos os .md
    md_files = sorted(BASE_DIR.rglob("*.md"))

    # Excluir o proprio script se houver .md
    md_files = [f for f in md_files if 'VERSAO DOC' not in str(f)]

    print(f"Total de arquivos .md encontrados: {len(md_files)}")
    print()

    success = 0
    errors = []

    for md_path in md_files:
        # Calcular caminho relativo
        rel = md_path.relative_to(BASE_DIR)
        docx_rel = rel.with_suffix('.docx')
        docx_path = OUTPUT_DIR / docx_rel

        print(f"  Convertendo: {rel.name}")
        print(f"  -> {docx_path.relative_to(BASE_DIR)}")

        try:
            convert_md_to_docx(md_path, docx_path)
            success += 1
            print(f"     [OK]")
        except Exception as e:
            errors.append((str(md_path), str(e)))
            print(f"     [ERRO] {e}")
        print()

    print("=" * 60)
    print(f"RESULTADO FINAL")
    print(f"  Arquivos .md encontrados: {len(md_files)}")
    print(f"  Arquivos .docx gerados:   {success}")
    print(f"  Erros:                    {len(errors)}")

    if errors:
        print("\nARQUIVOS COM ERRO:")
        for path, err in errors:
            print(f"  {path}: {err}")

    # Validacao final
    print("\nVALIDACAO:")
    all_ok = True
    for md_path in md_files:
        rel = md_path.relative_to(BASE_DIR)
        docx_rel = rel.with_suffix('.docx')
        docx_path = OUTPUT_DIR / docx_rel
        status = "OK" if docx_path.exists() else "FALTANDO"
        print(f"  [{status}] {rel.name} -> {docx_rel.name}")
        if status == "FALTANDO":
            all_ok = False

    if all_ok:
        print("\nTodos os arquivos foram convertidos com sucesso!")
    else:
        print("\nATENCAO: Alguns arquivos nao foram convertidos!")


if __name__ == "__main__":
    main()
